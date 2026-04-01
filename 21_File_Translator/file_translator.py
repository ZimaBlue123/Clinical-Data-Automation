# -*- coding: utf-8 -*-
"""
多格式文档翻译模块（Excel/CSV/Word/PDF）
优化项：
1) translators -> DeepL -> LibreTranslate 回退链
2) 英中/中英双向
3) PDF 可选仅导出双语文本层
4) Word 增强覆盖（页眉页脚/文本框/脚注）
5) 并发处理 + JSON 持久化缓存
6) 术语词典优先替换
"""

from __future__ import annotations

import argparse
import json
import os
import platform
import re
import shutil
import threading
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from copy import copy
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Protocol

import pandas as pd
import requests
from openpyxl import load_workbook
from openpyxl.cell.cell import Cell

BASE = Path(__file__).resolve().parent
INPUT_DIR = BASE / "input"
OUTPUT_DIR = BASE / "output"

DEFAULT_DEEPL_API_BASE = "https://api-free.deepl.com/v2"
DEFAULT_LIBRE_API_BASE = "https://libretranslate.com"
DEFAULT_TRANSLATE_COLUMNS = ["Term", "SOC", "Comment", "PT Name", "SOC Name", "English_Term"]

# COM 常量，避免魔法数字
XL_PASTE_FORMATS = -4122
WD_FORMAT_XML_DOCUMENT = 16


@dataclass(frozen=True)
class TransConfig:
    input_path: Path
    output_dir: Path
    direction: str
    columns: list[str]
    provider: str
    ts_engine: str
    ts_sleep: float
    deepl_key: str
    deepl_api_base: str
    libre_key: str
    libre_api_base: str
    engine: str
    pdf_mode: str
    word_include_textboxes: bool
    word_include_footnotes: bool
    max_workers: int
    cache_file: Path
    no_cache: bool
    glossary_path: Path | None


def _load_local_env(env_path: Path) -> None:
    if not env_path.exists():
        return
    for raw in env_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        k, v = line.split("=", 1)
        key = k.strip()
        val = v.strip().strip("'").strip('"')
        if key and key not in os.environ:
            os.environ[key] = val


def _mask_secret(s: str, head: int = 4, tail: int = 4) -> str:
    s = (s or "").strip()
    if len(s) <= head + tail:
        return "*" * len(s)
    return f"{s[:head]}{'*' * (len(s) - head - tail)}{s[-tail:]}"


def _suffix_by_direction(direction: str) -> str:
    return "CN" if direction == "en2zh" else "EN"


def _lang_pair(direction: str) -> tuple[str, str]:
    return ("en", "zh") if direction == "en2zh" else ("zh", "en")


def _normalize_str(v) -> str:
    return "" if v is None else str(v).strip()


class TranslationMemory:
    def __init__(self, path: Path | None, enabled: bool = True) -> None:
        self.path = path
        self.enabled = enabled
        self.lock = threading.Lock()
        self.data: dict[str, str] = {}
        if enabled and path and path.exists():
            try:
                self.data = json.loads(path.read_text(encoding="utf-8"))
            except Exception:
                self.data = {}

    def get(self, key: str) -> str | None:
        if not self.enabled:
            return None
        with self.lock:
            return self.data.get(key)

    def set(self, key: str, value: str) -> None:
        if not self.enabled:
            return
        with self.lock:
            self.data[key] = value

    def save(self) -> None:
        if not self.enabled or not self.path:
            return
        self.path.parent.mkdir(parents=True, exist_ok=True)
        with self.lock:
            tmp_path = self.path.with_suffix(f"{self.path.suffix}.tmp")
            tmp_path.write_text(json.dumps(self.data, ensure_ascii=False, indent=2), encoding="utf-8")
            os.replace(tmp_path, self.path)


def _load_glossary(path: Path | None) -> dict[str, str]:
    if not path or not path.exists():
        return {}
    if path.suffix.lower() == ".json":
        try:
            obj = json.loads(path.read_text(encoding="utf-8"))
            return {str(k): str(v) for k, v in obj.items()}
        except Exception:
            return {}
    out: dict[str, str] = {}
    for ln in path.read_text(encoding="utf-8").splitlines():
        if not ln.strip() or ln.strip().startswith("#"):
            continue
        if "\t" in ln:
            k, v = ln.split("\t", 1)
        elif "," in ln:
            k, v = ln.split(",", 1)
        else:
            continue
        out[k.strip()] = v.strip()
    return out


class TranslatorProtocol(Protocol):
    def batch_translate(self, texts: Iterable[object]) -> list[str]:
        ...

    def self_test(self) -> tuple[bool, str]:
        ...


class TranslatorsFreeTranslator:
    def __init__(self, engine: str, direction: str, sleep_s: float = 0.5) -> None:
        self.engine = (engine or "bing").strip().lower()
        self.sleep_s = max(0.0, sleep_s)
        self.src, self.dst = _lang_pair(direction)
        self.last_error = ""
        self._ts = None
        try:
            import translators as ts  # type: ignore

            self._ts = ts
        except Exception as e:  # noqa: BLE001
            self.last_error = str(e)

    def batch_translate(self, texts: Iterable[object]) -> list[str]:
        out: list[str] = []
        for x in texts:
            s = _normalize_str(x)
            if not s:
                out.append("")
                continue
            if self._ts is None:
                out.append(s)
                continue
            try:
                tr = str(self._ts.translate_text(s[:4800], translator=self.engine, from_language=self.src, to_language=self.dst)).strip()
                out.append(tr if tr else s)
                time.sleep(self.sleep_s)
            except Exception as e:  # noqa: BLE001
                self.last_error = str(e)
                out.append(s)
        return out

    def self_test(self) -> tuple[bool, str]:
        if self._ts is None:
            return False, "未安装 translators 包"
        sample = "Serious adverse event" if self.src == "en" else "严重不良事件"
        r = self.batch_translate([sample])[0]
        return (True, f"自检成功: {sample} -> {r}") if r and r != sample else (False, "自检失败")


class DeepLTranslator:
    def __init__(self, api_key: str, api_base: str, direction: str, sleep_s: float = 0.12) -> None:
        self.api_key = api_key.strip()
        self.api_base = api_base.rstrip("/")
        self.sleep_s = sleep_s
        self.src, self.dst = _lang_pair(direction)
        self.last_error = ""
        self.session = requests.Session()

    def batch_translate(self, texts: Iterable[object]) -> list[str]:
        src_lang = "EN" if self.src == "en" else "ZH"
        dst_lang = "ZH" if self.dst == "zh" else "EN"
        out: list[str] = []
        for x in texts:
            s = _normalize_str(x)
            if not s:
                out.append("")
                continue
            try:
                r = self.session.post(
                    f"{self.api_base}/translate",
                    data={
                        "auth_key": self.api_key,
                        "text": s[:4800],
                        "source_lang": src_lang,
                        "target_lang": dst_lang,
                        "preserve_formatting": "1",
                    },
                    timeout=60,
                )
                r.raise_for_status()
                tr = str(r.json().get("translations", [{}])[0].get("text", "")).strip()
                out.append(tr if tr else s)
                time.sleep(self.sleep_s)
            except Exception as e:  # noqa: BLE001
                self.last_error = str(e)
                out.append(s)
        return out

    def self_test(self) -> tuple[bool, str]:
        sample = "Serious adverse event" if self.src == "en" else "严重不良事件"
        r = self.batch_translate([sample])[0]
        return (True, f"自检成功: {sample} -> {r}") if r and r != sample else (False, f"自检失败: {self.last_error or '返回原文'}")


class LibreTranslateTranslator:
    def __init__(self, api_base: str, direction: str, api_key: str = "") -> None:
        self.api_base = api_base.rstrip("/")
        self.api_key = api_key.strip()
        self.src, self.dst = _lang_pair(direction)
        self.last_error = ""
        self.session = requests.Session()

    def batch_translate(self, texts: Iterable[object]) -> list[str]:
        out: list[str] = []
        for x in texts:
            s = _normalize_str(x)
            if not s:
                out.append("")
                continue
            payload = {"q": s[:4800], "source": self.src, "target": self.dst, "format": "text"}
            if self.api_key:
                payload["api_key"] = self.api_key
            try:
                r = self.session.post(f"{self.api_base}/translate", data=payload, timeout=60)
                r.raise_for_status()
                tr = str(r.json().get("translatedText", "")).strip()
                out.append(tr if tr else s)
            except Exception as e:  # noqa: BLE001
                self.last_error = str(e)
                out.append(s)
        return out

    def self_test(self) -> tuple[bool, str]:
        sample = "Serious adverse event" if self.src == "en" else "严重不良事件"
        r = self.batch_translate([sample])[0]
        return (True, f"自检成功: {sample} -> {r}") if r and r != sample else (False, f"自检失败: {self.last_error or '返回原文'}")


class FallbackTranslator:
    def __init__(self, candidates: list[tuple[str, TranslatorProtocol]]) -> None:
        self.candidates = candidates
        self.active_name = ""
        self.active: TranslatorProtocol | None = None

    def _ensure_active(self) -> tuple[bool, str]:
        if self.active is not None:
            return True, self.active_name
        for name, tr in self.candidates:
            ok, _ = tr.self_test()
            if ok:
                self.active_name = name
                self.active = tr
                return True, name
        return False, "all_failed"

    def batch_translate(self, texts: Iterable[object]) -> list[str]:
        ok, _ = self._ensure_active()
        if not ok or self.active is None:
            return [_normalize_str(x) for x in texts]
        return self.active.batch_translate(texts)

    def self_test(self) -> tuple[bool, str]:
        ok, name = self._ensure_active()
        return (True, f"自动回退已选中引擎: {name}") if ok else (False, "自动回退失败")


class SmartTranslator:
    def __init__(self, base: TranslatorProtocol, direction: str, memory: TranslationMemory, glossary: dict[str, str]) -> None:
        self.base = base
        self.direction = direction
        self.memory = memory
        self.glossary = glossary

    def _apply_glossary(self, text: str) -> tuple[str, dict[str, str]]:
        if not self.glossary:
            return text, {}
        tokens: dict[str, str] = {}
        protected = text
        for i, (src, dst) in enumerate(sorted(self.glossary.items(), key=lambda x: len(x[0]), reverse=True)):
            if src and src in protected:
                tok = f"__TERM_{i}__"
                protected = protected.replace(src, tok)
                tokens[tok] = dst
        return protected, tokens

    @staticmethod
    def _restore_glossary(text: str, tokens: dict[str, str]) -> str:
        out = text
        for tok, dst in tokens.items():
            out = out.replace(tok, dst)
        return out

    def batch_translate(self, texts: Iterable[object]) -> list[str]:
        raw = [_normalize_str(x) for x in texts]
        result: list[str] = []
        pending_idx: list[int] = []
        pending_text: list[str] = []
        pending_tokens: list[dict[str, str]] = []

        for i, s in enumerate(raw):
            if not s:
                result.append("")
                continue
            if s in self.glossary:
                result.append(self.glossary[s])
                continue
            key = f"{self.direction}|{s}"
            cached = self.memory.get(key)
            if cached is not None:
                result.append(cached)
                continue
            p, tks = self._apply_glossary(s)
            pending_idx.append(i)
            pending_text.append(p)
            pending_tokens.append(tks)
            result.append("")

        if pending_text:
            translated = self.base.batch_translate(pending_text)
            for i, idx in enumerate(pending_idx):
                src = raw[idx]
                dst = self._restore_glossary(translated[i], pending_tokens[i]) if i < len(translated) else src
                if not dst:
                    dst = src
                result[idx] = dst
                self.memory.set(f"{self.direction}|{src}", dst)
        return result

    def self_test(self) -> tuple[bool, str]:
        return self.base.self_test()


def _try_import_win32():
    try:
        import win32com.client as win32  # type: ignore
        return win32
    except Exception:  # noqa: BLE001
        return None


def _ensure_com_thread() -> None:
    try:
        import pythoncom  # type: ignore

        pythoncom.CoInitialize()
    except Exception:
        pass


def _copy_col_style(ws, src_col_idx: int, dst_col_idx: int, max_row: int) -> None:
    src_letter = ws.cell(1, src_col_idx).column_letter
    dst_letter = ws.cell(1, dst_col_idx).column_letter
    ws.column_dimensions[dst_letter].width = ws.column_dimensions[src_letter].width
    for r in range(1, max_row + 1):
        s: Cell = ws.cell(r, src_col_idx)
        d: Cell = ws.cell(r, dst_col_idx)
        d._style = copy(s._style)  # noqa: SLF001
        d.number_format = s.number_format
        d.alignment = copy(s.alignment)


def _iter_docx_runs(doc):
    for p in doc.paragraphs:
        for run in p.runs:
            yield run
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        yield run
    for sec in doc.sections:
        for hf in [sec.header, sec.footer]:
            for p in hf.paragraphs:
                for run in p.runs:
                    yield run
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                yield run


def _translate_word_com_extras(docx_path: Path, translator: TranslatorProtocol, include_textboxes: bool, include_footnotes: bool) -> None:
    _ensure_com_thread()
    win32 = _try_import_win32()
    if win32 is None:
        return
    word = win32.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False
    doc = None
    try:
        doc = word.Documents.Open(str(docx_path.resolve()))
        if include_textboxes:
            for shape in doc.Shapes:
                try:
                    if shape.TextFrame.HasText:
                        txt = _normalize_str(shape.TextFrame.TextRange.Text)
                        if txt:
                            shape.TextFrame.TextRange.Text = translator.batch_translate([txt])[0]
                except Exception:
                    pass
        if include_footnotes:
            try:
                for fn in doc.Footnotes:
                    txt = _normalize_str(fn.Range.Text)
                    if txt:
                        fn.Range.Text = translator.batch_translate([txt])[0]
            except Exception:
                pass
        doc.Save()
    finally:
        if doc is not None:
            doc.Close(False)
        word.Quit()


def translate_xlsx_openpyxl(input_path: Path, output_path: Path, translator: TranslatorProtocol, columns: list[str], suffix: str) -> None:
    wb = load_workbook(input_path)
    for ws in wb.worksheets:
        header_map: dict[str, int] = {}
        for c in range(1, ws.max_column + 1):
            key = _normalize_str(ws.cell(1, c).value)
            if key:
                header_map[key] = c
        hit_cols = [c for c in columns if c in header_map]
        if not hit_cols:
            continue
        append_start = ws.max_column + 1
        for i, src_name in enumerate(hit_cols):
            src_idx = header_map[src_name]
            dst_idx = append_start + i
            ws.cell(1, dst_idx).value = f"{src_name}_{suffix}"
            _copy_col_style(ws, src_idx, dst_idx, ws.max_row)
            src_values = [ws.cell(r, src_idx).value for r in range(2, ws.max_row + 1)]
            out_values = translator.batch_translate(src_values)
            for row_idx, val in enumerate(out_values, start=2):
                ws.cell(row_idx, dst_idx).value = val
    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(output_path)


def translate_xlsx_com(input_path: Path, output_path: Path, translator: TranslatorProtocol, columns: list[str], suffix: str) -> None:
    _ensure_com_thread()
    win32 = _try_import_win32()
    if win32 is None:
        raise RuntimeError("win32com 不可用")
    shutil.copy2(input_path, output_path)
    excel = win32.DispatchEx("Excel.Application")
    excel.Visible = False
    excel.DisplayAlerts = False
    wb = None
    try:
        wb = excel.Workbooks.Open(str(output_path.resolve()))
        for ws in wb.Worksheets:
            used = ws.UsedRange
            max_row = int(used.Rows.Count)
            max_col = int(used.Columns.Count)
            header_map: dict[str, int] = {}
            for c in range(1, max_col + 1):
                key = _normalize_str(ws.Cells(1, c).Value)
                if key:
                    header_map[key] = c
            hit_cols = [c for c in columns if c in header_map]
            if not hit_cols:
                continue
            append_start = max_col + 1
            for i, src_name in enumerate(hit_cols):
                src_idx = header_map[src_name]
                dst_idx = append_start + i
                ws.Cells(1, dst_idx).Value = f"{src_name}_{suffix}"
                ws.Columns(dst_idx).ColumnWidth = ws.Columns(src_idx).ColumnWidth
                ws.Columns(src_idx).Copy()
                ws.Columns(dst_idx).PasteSpecial(Paste=XL_PASTE_FORMATS)
                src_values = [ws.Cells(r, src_idx).Value for r in range(2, max_row + 1)]
                out_values = translator.batch_translate(src_values)
                for row_idx, val in enumerate(out_values, start=2):
                    ws.Cells(row_idx, dst_idx).Value = val
        wb.Save()
    finally:
        if wb is not None:
            wb.Close(SaveChanges=True)
        excel.Quit()


def translate_csv(input_path: Path, output_path: Path, translator: TranslatorProtocol, columns: list[str], suffix: str) -> None:
    if "29.0 IME List" in input_path.name:
        df = pd.read_csv(input_path, skiprows=11, encoding="utf-8-sig")
    else:
        df = pd.read_csv(input_path, encoding="utf-8-sig")
    for c in columns:
        if c in df.columns:
            df[f"{c}_{suffix}"] = translator.batch_translate(df[c].tolist())
    output_path.parent.mkdir(parents=True, exist_ok=True)
    df.to_excel(output_path, index=False, engine="openpyxl")


def translate_docx(input_path: Path, output_path: Path, translator: TranslatorProtocol, include_textboxes: bool, include_footnotes: bool) -> None:
    from docx import Document

    doc = Document(input_path)
    for run in _iter_docx_runs(doc):
        txt = _normalize_str(run.text)
        if txt:
            run.text = translator.batch_translate([txt])[0]
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    if include_textboxes or include_footnotes:
        _translate_word_com_extras(output_path, translator, include_textboxes, include_footnotes)


def translate_doc_via_com(input_path: Path, output_path: Path, translator: TranslatorProtocol, include_textboxes: bool, include_footnotes: bool) -> None:
    _ensure_com_thread()
    win32 = _try_import_win32()
    if win32 is None:
        raise RuntimeError("win32com 不可用，无法处理 .doc")
    tmp_docx = output_path.with_suffix(".docx")
    word = win32.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False
    doc = None
    try:
        doc = word.Documents.Open(str(input_path.resolve()))
        doc.SaveAs(str(tmp_docx.resolve()), FileFormat=WD_FORMAT_XML_DOCUMENT)
    finally:
        if doc is not None:
            doc.Close(False)
        word.Quit()
    translate_docx(tmp_docx, output_path.with_suffix(".docx"), translator, include_textboxes, include_footnotes)


def _resolve_cjk_fontfile() -> str | None:
    """PyMuPDF 叠字中文需 fontfile；纯 fontname='cjk' 在部分环境会报 need font file or buffer。"""
    env = os.environ.get("PDF_CJK_FONT", "").strip()
    if env and Path(env).is_file():
        return env
    if platform.system().lower() == "windows":
        windir = Path(os.environ.get("WINDIR", r"C:\Windows"))
        for name in ("msyh.ttc", "msyhbd.ttc", "simsun.ttc", "simhei.ttf"):
            p = windir / "Fonts" / name
            if p.is_file():
                return str(p)
    return None


def translate_pdf_overlay(input_path: Path, output_path: Path, translator: TranslatorProtocol) -> None:
    import fitz

    cjk_font = _resolve_cjk_fontfile()
    if not cjk_font:
        raise RuntimeError(
            "未找到 CJK 字体文件，无法 PDF overlay 中文。请设置环境变量 PDF_CJK_FONT 指向 .ttf/.ttc，"
            "或使用 --pdf-mode bilingual-text-layer 仅导出双语对照 txt。"
        )

    doc = fitz.open(input_path)
    for page in doc:
        info = page.get_text("dict")
        # 批量翻译：把本页所有 span 文本收集起来一次性送翻译引擎，显著减少网络调用次数。
        spans: list[tuple[fitz.Rect, float, str]] = []
        texts: list[str] = []
        for block in info.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    txt = _normalize_str(span.get("text"))
                    if not txt:
                        continue
                    # bbox 可能比较紧凑，直接用 span bbox + 较大字体容易被裁剪导致“白底但无字”。
                    # 策略：先插入翻译文字（不擦除原文），只有当 textbox 返回值提示“有余量/可容纳”
                    # 时才在 bbox 内擦除原文，避免大面积空白。
                    rect = fitz.Rect(span["bbox"])
                    pad = 0.2
                    rect = fitz.Rect(rect.x0 - pad, rect.y0 - pad, rect.x1 + pad, rect.y1 + pad)
                    rect_h = max(1e-6, rect.y1 - rect.y0)
                    fontsize_raw = float(span.get("size", 10) or 10)
                    fontsize = max(4.0, min(fontsize_raw, rect_h * 0.9))
                    spans.append((rect, fontsize, txt))
                    texts.append(txt)

        if not texts:
            continue

        translated = translator.batch_translate(texts)
        for (rect, fontsize, src), tr in zip(spans, translated):
            if not tr or tr == src:
                continue
            lineheight = fontsize
            unused_area = page.insert_textbox(
                rect,
                tr,
                fontsize=fontsize,
                fontfile=cjk_font,
                fontname="f0",
                color=(0, 0, 0),
                align=0,
                # fill_opacity=0 会导致文字本身透明
                fill_opacity=1,
                overlay=True,
            )
            # PyMuPDF：返回值为“未使用/剩余”的区域面积（可能为负代表溢出/裁剪）。
            if unused_area >= 0:
                page.draw_rect(
                    rect,
                    color=(1, 1, 1),
                    fill=(1, 1, 1),
                    width=0,
                    overlay=False,  # 不覆盖刚插入的翻译文字
                    stroke_opacity=0,
                    fill_opacity=1,
                )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    doc.close()


def translate_pdf_bilingual_layer(input_path: Path, output_path: Path, translator: TranslatorProtocol) -> None:
    import fitz

    shutil.copy2(input_path, output_path)
    doc = fitz.open(input_path)
    lines: list[str] = []
    for pno, page in enumerate(doc, start=1):
        text = page.get_text("text")
        src_lines = [x.strip() for x in text.splitlines() if x.strip()]
        if not src_lines:
            continue
        tr_lines = translator.batch_translate(src_lines)
        lines.append(f"# Page {pno}")
        for s, t in zip(src_lines, tr_lines):
            lines.append(f"SRC: {s}")
            lines.append(f"TR : {t}")
            lines.append("")
    doc.close()
    sidecar = output_path.with_suffix(".bilingual.txt")
    sidecar.write_text("\n".join(lines), encoding="utf-8")


def build_translator(cfg: TransConfig, memory: TranslationMemory, glossary: dict[str, str]) -> SmartTranslator:
    if cfg.provider == "tsfree":
        base: TranslatorProtocol = TranslatorsFreeTranslator(cfg.ts_engine, cfg.direction, cfg.ts_sleep)
    elif cfg.provider == "deepl":
        base = DeepLTranslator(cfg.deepl_key, cfg.deepl_api_base, cfg.direction)
    elif cfg.provider == "libre":
        base = LibreTranslateTranslator(cfg.libre_api_base, cfg.direction, cfg.libre_key or "")
    else:
        cands: list[tuple[str, TranslatorProtocol]] = [("tsfree", TranslatorsFreeTranslator(cfg.ts_engine, cfg.direction, cfg.ts_sleep))]
        if cfg.deepl_key:
            cands.append(("deepl", DeepLTranslator(cfg.deepl_key, cfg.deepl_api_base, cfg.direction)))
        cands.append(("libre", LibreTranslateTranslator(cfg.libre_api_base, cfg.direction, cfg.libre_key or "")))
        base = FallbackTranslator(cands)
    return SmartTranslator(base, cfg.direction, memory, glossary)


def process_one_file(input_path: Path, cfg: TransConfig, translator: SmartTranslator) -> Path:
    suffix = _suffix_by_direction(cfg.direction)
    src_suffix = input_path.suffix.lower()
    out_ext = ".xlsx" if src_suffix in {".xlsx", ".csv"} else (".docx" if src_suffix == ".doc" else src_suffix)
    out = cfg.output_dir / f"{input_path.stem}_{cfg.direction}{out_ext}"

    if src_suffix == ".xlsx":
        if cfg.engine == "com":
            translate_xlsx_com(input_path, out, translator, cfg.columns, suffix)
        else:
            translate_xlsx_openpyxl(input_path, out, translator, cfg.columns, suffix)
    elif src_suffix == ".csv":
        translate_csv(input_path, out, translator, cfg.columns, suffix)
    elif src_suffix == ".docx":
        translate_docx(input_path, out, translator, cfg.word_include_textboxes, cfg.word_include_footnotes)
    elif src_suffix == ".doc":
        translate_doc_via_com(input_path, out, translator, cfg.word_include_textboxes, cfg.word_include_footnotes)
    elif src_suffix == ".pdf":
        if cfg.pdf_mode == "bilingual-text-layer":
            translate_pdf_bilingual_layer(input_path, out, translator)
        else:
            translate_pdf_overlay(input_path, out, translator)
    else:
        raise ValueError(f"不支持的文件类型: {input_path.name}")
    return out


def main() -> None:
    _load_local_env(BASE / ".env")
    parser = argparse.ArgumentParser(description="21 模块：Excel/CSV/Word/PDF 翻译（免费优先）")
    parser.add_argument("--input", "-i", default=str(INPUT_DIR), help="输入文件或目录")
    parser.add_argument("--output", "-o", default=str(OUTPUT_DIR), help="输出目录")
    parser.add_argument("--direction", choices=["en2zh", "zh2en"], default=os.getenv("TRANSLATION_DIRECTION", "en2zh"), help="翻译方向")
    parser.add_argument("--columns", default=",".join(DEFAULT_TRANSLATE_COLUMNS), help="Excel/CSV 翻译列（逗号分隔）")
    parser.add_argument("--provider", choices=["auto", "tsfree", "deepl", "libre"], default=os.getenv("TRANSLATOR_PROVIDER", "auto"), help="翻译引擎")
    parser.add_argument("--ts-engine", default=os.getenv("TS_TRANSLATOR_ENGINE", "bing"), help="translators 引擎（默认 bing）")
    parser.add_argument("--ts-sleep", type=float, default=float(os.getenv("TS_SLEEP_SECONDS", "0.5")), help="translators 请求休眠秒数")
    parser.add_argument("--deepl-key", default=os.getenv("DEEPL_API_KEY"), help="DeepL API Key")
    parser.add_argument("--deepl-api-base", default=os.getenv("DEEPL_API_BASE") or DEFAULT_DEEPL_API_BASE, help="DeepL API Base")
    parser.add_argument("--libre-key", default=os.getenv("LIBRETRANSLATE_API_KEY"), help="LibreTranslate API Key（可选）")
    parser.add_argument("--libre-api-base", default=os.getenv("LIBRETRANSLATE_API_BASE") or DEFAULT_LIBRE_API_BASE, help="LibreTranslate API Base")
    parser.add_argument("--self-test", action="store_true", help="仅执行连通测试")
    parser.add_argument("--skip-preflight", action="store_true", help="跳过处理前自检")
    parser.add_argument("--engine", choices=["auto", "openpyxl", "com"], default="auto", help="xlsx 写回引擎")
    parser.add_argument("--pdf-mode", choices=["overlay", "bilingual-text-layer"], default=os.getenv("PDF_TRANSLATE_MODE", "overlay"), help="PDF 翻译模式")
    parser.add_argument("--word-include-textboxes", action="store_true", default=True, help="Word 启用文本框翻译（COM 可用时）")
    parser.add_argument("--word-include-footnotes", action="store_true", default=True, help="Word 启用脚注翻译（COM 可用时）")
    parser.add_argument("--max-workers", type=int, default=int(os.getenv("MAX_WORKERS", "1")), help="并发文件数")
    parser.add_argument("--cache-file", default=os.getenv("TRANSLATION_CACHE_FILE") or str(OUTPUT_DIR / "translation_cache.json"), help="持久化缓存文件")
    parser.add_argument("--no-cache", action="store_true", help="禁用持久化缓存")
    parser.add_argument("--glossary", default=os.getenv("GLOSSARY_FILE"), help="术语词典文件（json/csv/tsv）")
    args = parser.parse_args()

    if args.provider == "deepl" and not args.deepl_key:
        raise ValueError("provider=deepl 时必须提供 DEEPL_API_KEY")

    engine = args.engine
    if engine == "auto":
        engine = "com" if platform.system().lower() == "windows" and _try_import_win32() is not None else "openpyxl"
    if args.max_workers > 1 and engine == "com":
        print("[WARN] 并发模式下 COM 引擎线程不安全，自动切换为 openpyxl。")
        engine = "openpyxl"

    cfg = TransConfig(
        input_path=Path(args.input),
        output_dir=Path(args.output),
        direction=args.direction,
        columns=[x.strip() for x in args.columns.split(",") if x.strip()],
        provider=args.provider,
        ts_engine=args.ts_engine,
        ts_sleep=args.ts_sleep,
        deepl_key=args.deepl_key or "",
        deepl_api_base=args.deepl_api_base,
        libre_key=args.libre_key or "",
        libre_api_base=args.libre_api_base,
        engine=engine,
        pdf_mode=args.pdf_mode,
        word_include_textboxes=args.word_include_textboxes,
        word_include_footnotes=args.word_include_footnotes,
        max_workers=max(1, args.max_workers),
        cache_file=Path(args.cache_file),
        no_cache=args.no_cache,
        glossary_path=Path(args.glossary) if args.glossary else None,
    )

    mem = TranslationMemory(cfg.cache_file, enabled=not cfg.no_cache)
    glossary = _load_glossary(cfg.glossary_path)
    translator = build_translator(cfg, mem, glossary)

    print(
        f"配置: provider={cfg.provider}, direction={cfg.direction}, pdf_mode={cfg.pdf_mode}, "
        f"workers={cfg.max_workers}, cache={'on' if not cfg.no_cache else 'off'}, glossary={len(glossary)}"
    )
    if cfg.provider == "deepl":
        print(f"DeepL key: {_mask_secret(cfg.deepl_key)}")

    if args.self_test:
        ok, msg = translator.self_test()
        print(msg)
        raise SystemExit(0 if ok else 2)
    if not args.skip_preflight:
        ok, msg = translator.self_test()
        print(msg)
        if not ok:
            raise SystemExit(2)

    cfg.output_dir.mkdir(parents=True, exist_ok=True)
    print(f"写回引擎: {cfg.engine}")

    exts = {".xlsx", ".csv", ".docx", ".doc", ".pdf"}
    if cfg.input_path.is_file():
        files = [cfg.input_path]
    elif cfg.input_path.is_dir():
        files = [p for p in cfg.input_path.iterdir() if p.is_file() and p.suffix.lower() in exts and not p.name.startswith("~$")]
    else:
        raise FileNotFoundError(f"输入不存在: {cfg.input_path}")
    if not files:
        print(f"未找到可处理文件: {cfg.input_path}")
        return

    print(f"共 {len(files)} 个文件待处理。")

    def _worker(f: Path) -> tuple[Path, str]:
        local_translator = build_translator(cfg, mem, glossary)
        out = process_one_file(f, cfg, local_translator)
        return f, str(out)

    if cfg.max_workers <= 1:
        for idx, f in enumerate(files, start=1):
            print(f"[{idx}/{len(files)}] 处理中: {f.name}")
            try:
                _, out = _worker(f)
                print(f"  -> 输出: {out}")
            except Exception as e:  # noqa: BLE001
                print(f"[ERROR] 处理失败 {f.name}: {e}")
    else:
        with ThreadPoolExecutor(max_workers=cfg.max_workers) as ex:
            fut_map = {ex.submit(_worker, f): f for f in files}
            done_count = 0
            for fut in as_completed(fut_map):
                f = fut_map[fut]
                done_count += 1
                try:
                    _, out = fut.result()
                    print(f"[{done_count}/{len(files)}] 完成: {f.name} -> {out}")
                except Exception as e:  # noqa: BLE001
                    print(f"[ERROR] 处理失败 {f.name}: {e}")

    mem.save()
    print("完成。")


if __name__ == "__main__":
    main()

